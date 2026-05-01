"""
Модуль приёма и обработки документов в RAG.
Поддерживает: PDF (текст + изображения с OCR через Gemini), DOCX, TXT, XLSX.
"""
import io
import os
import base64
import logging
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from openpyxl import load_workbook

from app.config import settings

logger = logging.getLogger(__name__)

# Настройка Gemini через ProxyAPI (только для LLM и OCR)
genai.configure(
    api_key=settings.GEMINI_API_KEY,
    transport="rest",
    client_options={"api_endpoint": settings.GEMINI_BASE_URL},
)

# Локальные embeddings (sentence-transformers, без API)
embeddings = HuggingFaceEmbeddings(
    model_name="all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True},
)

# ChromaDB
vectorstore = Chroma(
    collection_name="baza_knowledge",
    embedding_function=embeddings,
    persist_directory=settings.CHROMA_PATH,
)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    separators=["\n\n", "\n", ".", "!", "?", ",", " "],
)


def _ocr_pdf_page_with_gemini(page_image_bytes: bytes) -> str:
    """OCR одной страницы PDF через Gemini Vision."""
    model = genai.GenerativeModel(settings.GEMINI_MODEL)
    img_b64 = base64.b64encode(page_image_bytes).decode("utf-8")
    response = model.generate_content(
        [
            {
                "inline_data": {
                    "mime_type": "image/png",
                    "data": img_b64,
                }
            },
            (
                "Извлеки весь текст с этого изображения страницы документа по недвижимости. "
                "Верни только текст, сохраняя структуру и форматирование. "
                "Если таблица — сохрани табличную структуру. "
                "Если текста нет — верни пустую строку."
            ),
        ]
    )
    return response.text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF. Если страница не содержит текста — применяет OCR через Gemini."""
    doc = fitz.open(file_path)
    full_text = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()

        if len(text) < 50:  # Страница без текста → OCR
            logger.info(f"Страница {page_num + 1} без текста — OCR через Gemini")
            # Рендерим страницу в PNG
            mat = fitz.Matrix(2, 2)  # 2x zoom для лучшего качества
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            text = _ocr_pdf_page_with_gemini(img_bytes)

        if text:
            full_text.append(f"[Страница {page_num + 1}]\n{text}")

    doc.close()
    return "\n\n".join(full_text)


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из DOCX."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Также таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def extract_text_from_xlsx(file_path: str) -> str:
    """Извлекает текст из Excel."""
    wb = load_workbook(file_path, read_only=True, data_only=True)
    rows = []
    for sheet in wb.worksheets:
        rows.append(f"[Лист: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None)
            if row_text.strip():
                rows.append(row_text)
    wb.close()
    return "\n".join(rows)


def extract_text_from_txt(file_path: str) -> str:
    """Читает текстовый файл."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str, file_type: str) -> str:
    """Универсальная функция извлечения текста по типу файла."""
    ft = file_type.lower()
    if ft == "pdf":
        return extract_text_from_pdf(file_path)
    elif ft in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    elif ft in ("xlsx", "xls"):
        return extract_text_from_xlsx(file_path)
    elif ft in ("txt", "md"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {file_type}")


def ingest_document(
    file_path: str,
    file_type: str,
    document_id: int,
    original_name: str,
    project_slug: Optional[str] = None,
    project_name: Optional[str] = None,
) -> int:
    """
    Обрабатывает документ: извлекает текст, разбивает на чанки, добавляет в ChromaDB.
    Возвращает количество добавленных чанков.
    """
    logger.info(f"Начинаю обработку: {original_name} (тип={file_type}, проект={project_slug})")

    raw_text = extract_text(file_path, file_type)

    if not raw_text.strip():
        raise ValueError("Документ не содержит читаемого текста")

    # Метаданные для фильтрации
    metadata = {
        "document_id": str(document_id),
        "filename": original_name,
        "file_type": file_type,
        "project_slug": project_slug or "global",
        "project_name": project_name or "Общее",
    }

    chunks = text_splitter.create_documents(
        [raw_text],
        metadatas=[metadata] * 1,  # will be applied per chunk
    )

    # Обновляем метаданные для каждого чанка
    for i, chunk in enumerate(chunks):
        chunk.metadata = {**metadata, "chunk_index": str(i)}

    vectorstore.add_documents(chunks)
    logger.info(f"Документ {original_name}: добавлено {len(chunks)} чанков")
    return len(chunks)


def delete_document_from_vectorstore(document_id: int) -> None:
    """Удаляет все чанки документа из ChromaDB по document_id."""
    vectorstore.delete(
        where={"document_id": str(document_id)}
    )
    logger.info(f"Удалены чанки документа {document_id} из векторного хранилища")
